from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path("/Users/choies/Documents/gbsw conn/학교_소통_제안_시스템_제품_기획서.docx")

# Design preset: narrative_proposal, with a Korean-font override for portability.
# Header pattern: editorial_cover.
PAGE_WIDTH = Inches(8.5)
PAGE_HEIGHT = Inches(11)
MARGIN = Inches(1)
CONTENT_WIDTH = Inches(6.5)

INK = "1F2933"
MUTED = "5B6770"
BLUE = "2E74B5"
BLUE_DARK = "1F4D78"
PALE_BLUE = "EAF2F8"
PALE_GRAY = "F4F6F9"
MID_GRAY = "D9E0E6"
WHITE = "FFFFFF"
GREEN = "2E7D5B"
AMBER = "9A6700"
RED = "A23B3B"
# Match the installed variable font's exact family name so Word and Pages do not
# substitute a visually similar fallback.
FONT = "Pretendard Variable"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_inches: float) -> None:
    width_twips = round(width_inches * 1440)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[float], *, indent_twips: int = 120) -> None:
    width_twips = [round(width * 1440) for width in widths]
    total_twips = sum(width_twips)
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:w"), str(total_twips))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_twips))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_grid = table._tbl.tblGrid
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in width_twips:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)


def keep_with_next(paragraph, value=True) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    node = p_pr.find(qn("w:keepNext"))
    if node is None:
        node = OxmlElement("w:keepNext")
        p_pr.append(node)
    node.set(qn("w:val"), "1" if value else "0")


def keep_lines(paragraph, value=True) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    node = p_pr.find(qn("w:keepLines"))
    if node is None:
        node = OxmlElement("w:keepLines")
        p_pr.append(node)
    node.set(qn("w:val"), "1" if value else "0")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_1, instr_text, fld_char_2])
    run.font.name = FONT
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(MUTED)


def set_east_asia_font(run, font_name=FONT) -> None:
    run.font.name = font_name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{key}"), font_name)
    lang = r_pr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        r_pr.append(lang)
    lang.set(qn("w:val"), "ko-KR")
    lang.set(qn("w:eastAsia"), "ko-KR")
    lang.set(qn("w:bidi"), "ko-KR")


def style_run(run, *, size=None, bold=None, color=None, italic=None) -> None:
    set_east_asia_font(run)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def configure_numbering(document: Document) -> tuple[int, int]:
    numbering = document.part.numbering_part.element

    def add_abstract(abstract_id: int, fmt: str, text: str, left: int, hanging: int) -> None:
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "space")
        p_pr = OxmlElement("w:pPr")
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(left))
        ind.set(qn("w:hanging"), str(hanging))
        p_pr.append(ind)
        lvl.extend([start, num_fmt, lvl_text, suff, p_pr])
        abstract.append(lvl)
        numbering.append(abstract)

    def add_num(num_id: int, abstract_id: int) -> None:
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_num_id = OxmlElement("w:abstractNumId")
        abstract_num_id.set(qn("w:val"), str(abstract_id))
        num.append(abstract_num_id)
        numbering.append(num)

    add_abstract(42, "bullet", "•", 540, 280)
    add_num(42, 42)
    add_abstract(43, "decimal", "%1.", 540, 280)
    add_num(43, 43)
    return 42, 43


def apply_num(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])


def add_hyperlink(paragraph, text: str, url: str, color=BLUE) -> None:
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), FONT)
    r_fonts.set(qn("w:hAnsi"), FONT)
    r_fonts.set(qn("w:eastAsia"), FONT)
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.extend([r_fonts, c, u])
    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def setup_document() -> tuple[Document, int, int]:
    doc = Document()
    section = doc.sections[0]
    section.page_width = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT
    section.top_margin = MARGIN
    section.bottom_margin = MARGIN
    section.left_margin = MARGIN
    section.right_margin = MARGIN
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.34

    for name, size, color, before, after in (
        ("Title", 28, INK, 0, 14),
        ("Subtitle", 13, MUTED, 0, 8),
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 11.5, BLUE_DARK, 8, 4),
    ):
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(10.5)
        style.font.color.rgb = RGBColor.from_string(INK)
        style.paragraph_format.left_indent = Inches(0.28)
        style.paragraph_format.first_line_indent = Inches(-0.16)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.2

    header = section.header
    hp = header.paragraphs[0]
    hp.text = "학교 소통 제안 시스템  ·  제품 기획서"
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style_run(hp.runs[0], size=8.5, color=MUTED)
    bottom = OxmlElement("w:pBdr")
    bdr = OxmlElement("w:bottom")
    bdr.set(qn("w:val"), "single")
    bdr.set(qn("w:sz"), "4")
    bdr.set(qn("w:space"), "3")
    bdr.set(qn("w:color"), MID_GRAY)
    bottom.append(bdr)
    hp._p.get_or_add_pPr().append(bottom)

    footer = section.footer
    fp = footer.paragraphs[0]
    add_page_number(fp)

    doc.core_properties.title = "학교 소통 제안 시스템 제품 기획서"
    doc.core_properties.subject = "경북소프트웨어마이스터고 학생자치 기반 제안·정보공개 플랫폼"
    doc.core_properties.author = "프로젝트 기획팀"
    doc.core_properties.keywords = "학생자치, 제안, 익명성, 학교 소통, 오픈 소스"
    bullet_num, decimal_num = configure_numbering(doc)
    return doc, bullet_num, decimal_num


def add_body(doc: Document, text: str, *, bold_prefix: str | None = None, align=None) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if bold_prefix and text.startswith(bold_prefix):
        first = p.add_run(bold_prefix)
        style_run(first, bold=True, color=INK)
        rest = p.add_run(text[len(bold_prefix):])
        style_run(rest)
    else:
        run = p.add_run(text)
        style_run(run)
    keep_lines(p)


def add_bullets(doc: Document, items: list[str], bullet_num: int) -> None:
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.28)
        p.paragraph_format.first_line_indent = Inches(-0.16)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.2
        run = p.add_run("•  " + item)
        style_run(run)


def add_numbers(doc: Document, items: list[str], decimal_num: int) -> None:
    for index, item in enumerate(items, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.34)
        p.paragraph_format.first_line_indent = Inches(-0.24)
        p.paragraph_format.space_after = Pt(5)
        run = p.add_run(f"{index}.  {item}")
        style_run(run)


def add_callout(doc: Document, label: str, text: str, *, color=BLUE, fill=PALE_BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_geometry(table, [6.5], indent_twips=180)
    set_cant_split(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_width(cell, 6.5)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=150, start=180, bottom=150, end=180)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(label + "  ")
    style_run(r1, bold=True, color=color, size=9.5)
    r2 = p.add_run(text)
    style_run(r2, color=INK, size=9.5)
    keep_lines(p)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_geometry(table, widths, indent_twips=120)

    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    set_cant_split(hdr)
    for i, header in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_width(cell, widths[i])
        set_cell_shading(cell, PALE_GRAY)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        style_run(run, bold=True, color=BLUE_DARK, size=9)
        keep_lines(p)

    for row in rows:
        cells = table.add_row().cells
        set_cant_split(table.rows[-1])
        for i, value in enumerate(row):
            cell = cells[i]
            set_cell_width(cell, widths[i])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.12
            run = p.add_run(value)
            style_run(run, size=8.7, color=INK)
            keep_lines(p)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def section_heading(doc: Document, number: str, title: str, lead: str | None = None, *, new_page=True) -> None:
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(3)
    run = kicker.add_run(number)
    style_run(run, bold=True, size=8.5, color=BLUE)
    keep_with_next(kicker)
    h = doc.add_paragraph(title, style="Heading 1")
    keep_with_next(h)
    if lead:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(14)
        r = p.add_run(lead)
        style_run(r, size=11.5, color=MUTED)
        keep_lines(p)


def add_subheading(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="Heading 2")
    keep_with_next(p)


def add_small_label(doc: Document, text: str, color=BLUE) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    style_run(r, bold=True, size=8.5, color=color)
    keep_with_next(p)


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(70)
    r = p.add_run("제품 기획서  ·  DRAFT 0.1")
    style_run(r, bold=True, size=9, color=BLUE)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = title.add_run("학교 소통 제안 시스템")
    style_run(r, bold=True, size=28, color=INK)

    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = subtitle.add_run("경북소프트웨어마이스터고등학교 학생자치 기반\n제안·정보공개 플랫폼")
    style_run(r, size=13, color=MUTED)

    line = doc.add_paragraph()
    line.paragraph_format.space_before = Pt(22)
    line.paragraph_format.space_after = Pt(26)
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "16")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), BLUE)
    p_bdr.append(bottom)
    line._p.get_or_add_pPr().append(p_bdr)

    statement = doc.add_paragraph()
    statement.paragraph_format.space_after = Pt(28)
    r = statement.add_run("학생의 생각이 사라지지 않고,\n학교의 답변이 끝까지 기록되는 구조를 만든다.")
    style_run(r, bold=True, size=16, color=BLUE_DARK)

    meta = doc.add_table(rows=4, cols=2)
    meta.autofit = False
    meta.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(meta, [1.15, 5.35], indent_twips=0)
    labels = ["문서 목적", "주요 독자", "작성 기준일", "문서 상태"]
    values = [
        "학교 구성원 간 합의와 MVP 개발 착수를 위한 기준 문서",
        "학생, 교사, 학교 관리자, 개발팀, 오픈 소스 기여자",
        "2026년 8월 2일",
        "학교 승인 및 세부 운영정책 확정 전 기획 초안",
    ]
    for i in range(4):
        set_cell_width(meta.cell(i, 0), 1.15)
        set_cell_width(meta.cell(i, 1), 5.35)
        set_cell_margins(meta.cell(i, 0), top=50, start=0, bottom=50, end=120)
        set_cell_margins(meta.cell(i, 1), top=50, start=0, bottom=50, end=0)
        p1 = meta.cell(i, 0).paragraphs[0]
        p2 = meta.cell(i, 1).paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        p2.paragraph_format.space_after = Pt(0)
        r1 = p1.add_run(labels[i])
        r2 = p2.add_run(values[i])
        style_run(r1, bold=True, size=8.5, color=MUTED)
        style_run(r2, size=8.5, color=INK)


def build() -> None:
    doc, bullet_num, decimal_num = setup_document()
    add_cover(doc)
    doc.add_page_break()

    # 0. Executive summary
    section_heading(
        doc,
        "00  EXECUTIVE SUMMARY",
        "한눈에 보는 기획",
        "이 문서는 코드를 어떻게 짤지보다, 왜 이 서비스를 만들고 누구의 권리를 어떤 절차로 보호할지를 먼저 설명한다.",
    )
    add_callout(
        doc,
        "한 문장 정의",
        "로그인한 학생이 익명 또는 실명으로 학교 개선안을 제안하고, 50명의 동의를 받은 안건에 학교가 공식 답변하며, 민감한 익명 신원은 정해진 3인의 만장일치 없이는 확인할 수 없는 교내 소통 플랫폼이다.",
    )
    add_subheading(doc, "이 서비스가 해결하려는 일")
    add_bullets(
        doc,
        [
            "학생의 개선 아이디어가 개인적인 불만이나 일회성 대화로 사라지는 문제를 줄인다.",
            "학생자치가 의견 수렴을 넘어 정식 안건 형성과 실행 추적까지 담당할 수 있게 한다.",
            "교사는 모든 글을 즉시 처리하지 않고, 충분한 공감을 얻은 안건에 공식적으로 답변한다.",
            "익명 표현의 자유와 악의적 비방 대응 사이에 명확한 절차와 기록을 둔다.",
            "학교 재정·행정 정보를 출처와 함께 공개해 학생이 학교 운영을 이해할 수 있게 한다.",
        ],
        bullet_num,
    )
    add_subheading(doc, "현재까지 확정된 핵심 결정")
    add_table(
        doc,
        ["영역", "확정 내용", "의미"],
        [
            ["접근", "모든 사용자는 자체 계정으로 로그인", "공개 익명 게시판이 아니라 책임 있는 교내 참여 공간"],
            ["제안", "익명·실명 중 작성자가 선택", "공개 화면의 익명성과 시스템 내부 책임성을 함께 보장"],
            ["정식 안건", "활성 학생 50명 동의", "교사 업무를 보호하면서 대표성 있는 제안을 선별"],
            ["신원 확인", "학생부장·학생회장·부회장 만장일치", "권한 집중을 막고 절차적 정당성을 확보"],
            ["실제 열람", "학생부장만 재인증 후 가능", "학생 임원은 판단하되 신원은 보지 않음"],
            ["기술", "Next.js + Spring Boot + MySQL, Redis 제외", "학교에서 배우는 기술과 장기 유지보수성을 우선"],
            ["공개 방식", "MVP 이후 Apache-2.0 오픈 소스 검토", "재사용과 기여를 허용하되 운영 데이터는 철저히 분리"],
        ],
        [1.0, 2.8, 2.7],
    )
    add_callout(doc, "중요", "이 문서는 학교의 공식 승인이나 운영 규정을 대신하지 않는다. 미확정 항목은 임의로 구현하지 않고, 담당자 합의 후 운영정책과 코드에 반영한다.", color=AMBER, fill="FFF7E6")

    # 1. Background
    section_heading(
        doc,
        "01  BACKGROUND",
        "왜 이 서비스를 만드는가",
        "학교는 지식을 배우는 곳인 동시에, 공동체의 문제를 발견하고 토론하며 책임 있게 바꾸는 법을 연습하는 곳이다.",
    )
    add_subheading(doc, "기획의 출발점")
    add_body(doc, "학생이 학교에서 배워야 하는 것은 교과 지식만이 아니다. 서로 다른 의견을 듣고, 공적인 문제를 제안하고, 다수의 공감을 모으고, 결정권자의 답변과 실행을 확인하는 경험도 민주 시민 교육의 일부다. 이 기획은 그 경험을 일상적인 학교 운영 안에 넣으려는 시도다.")
    add_body(doc, "현재도 학생은 담임교사, 학생회, 설문, 대면 건의 등 여러 방식으로 의견을 낼 수 있다. 그러나 어떤 의견이 접수됐는지, 얼마나 공감을 얻었는지, 학교가 무엇을 검토했고 왜 그렇게 결정했는지가 한 흐름으로 남지 않는 경우가 많다. 그 결과 학생은 ‘말해도 달라지지 않는다’고 느끼고, 교사는 반복적이거나 대표성이 불분명한 요구를 개별적으로 대응해야 한다.")
    add_subheading(doc, "우리가 만들고 싶은 변화")
    add_bullets(
        doc,
        [
            "학생은 수업, 기숙사, 시설, 학생자치, 정보공개 등 학교생활의 개선안을 안전하게 제안한다.",
            "다른 학생은 제안을 읽고 한 사람당 한 번만 동의한다.",
            "50명의 동의를 얻은 제안은 정식 안건이 되고, 학교는 공식 답변과 후속 상태를 남긴다.",
            "채택되지 않은 안건에도 결정 사유가 남아, 다음 논의가 같은 지점에서 반복되지 않게 한다.",
            "채택된 안건은 ‘검토 중’에서 끝나지 않고 실행 중·완료까지 추적한다.",
        ],
        bullet_num,
    )
    add_callout(doc, "제품 비전", "학생의 목소리를 크게 만드는 데서 멈추지 않고, 학교가 듣고 판단하고 실행한 과정을 신뢰할 수 있는 기록으로 남긴다.")

    # 2. Principles
    section_heading(
        doc,
        "02  PRINCIPLES",
        "제품과 운영의 원칙",
        "기능이 충돌할 때 무엇을 우선할지 정해 두어야 서비스가 사람에 따라 흔들리지 않는다.",
    )
    principles = [
        ("참여", "학생이 이해하기 쉬운 언어와 짧은 흐름으로 의견을 낼 수 있어야 한다."),
        ("책임", "로그인은 필수이며, 시스템은 남용 대응에 필요한 최소한의 연결 정보를 보호해 보관한다."),
        ("표현의 자유", "욕설·비판을 이유로 작성 단계에서 AI나 금칙어가 글을 자동 차단하지 않는다."),
        ("절차적 보호", "신고, 공개 제한, 신원 확인, 실제 열람을 서로 다른 행위로 분리한다."),
        ("최소 권한", "누구도 직책 하나만으로 익명 신원을 마음대로 볼 수 없다. 슈퍼 어드민도 예외가 아니다."),
        ("투명성", "안건의 상태, 공식 답변, 결정 사유, 후속 계획을 학생에게 이해 가능한 형태로 공개한다."),
        ("운영 가능성", "교사가 모든 제안을 즉시 처리하지 않도록 50명 동의로 정식 안건을 선별한다."),
        ("검증 가능성", "권한 변경, 심의, 신원 열람, 상태 변경 등 민감한 행위는 감사 기록으로 남긴다."),
        ("Less, but better", "화면과 기능을 과장하지 않고 핵심 정보와 행동만 또렷하게 남긴다."),
    ]
    for title, text in principles:
        add_small_label(doc, title)
        add_body(doc, text)
    add_callout(doc, "원칙 적용 예", "‘욕설이 있을 수 있으니 작성 전에 막자’가 아니라, 우선 표현을 허용하고 신고와 합의된 심의 절차로 사후 대응한다. 동시에 원문 HTML 실행은 허용하지 않아 보안 공격은 차단한다.", color=GREEN, fill="EDF7F2")

    # 3. Stakeholders
    section_heading(
        doc,
        "03  PEOPLE & GOVERNANCE",
        "누가 무엇을 할 수 있는가",
        "역할은 서열이 아니라 책임의 묶음으로 설계한다. 학생회 임원도 제안과 동의에서는 일반 학생과 같은 한 사람이다.",
    )
    add_table(
        doc,
        ["사용자", "주요 행동", "제한과 보호"],
        [
            ["일반 학생", "계정 활성화, 제안 작성, 동의·철회, 신고, 진행 확인", "다른 사람의 신원·동의자 명단·내부 식별자 열람 불가"],
            ["학생회장·부회장", "일반 학생 기능 + 지정 사건의 공개 제한·신원 확인 심의", "심의 중에도 익명 작성자의 실제 이름·학번은 열람 불가"],
            ["일반 교사", "50명 이상 정식 안건 조회, 공식 답변, 처리 상태 관리", "50명 미만 일반 제안은 원칙적으로 조회 불가"],
            ["학생부장", "교사 기능 + 지정 사건 심의, 만장일치 후 신원 열람", "최근 재인증과 사유 기록이 있어야 실제 신원 열람 가능"],
            ["슈퍼 어드민", "계정·역할·보직·운영 설정 관리", "직책만으로 익명 신원 열람 불가, 모든 민감 변경 감사"],
        ],
        [1.15, 2.65, 2.70],
    )
    add_subheading(doc, "임기 기반 보직")
    add_body(doc, "학생부장, 전교 학생회장, 전교 학생부회장은 계정 역할과 별개인 임기 기반 보직이다. 슈퍼 어드민이 시작·종료 시각과 임명 사유를 포함해 지정하며, 다음 학기나 학년도에 새 담당자로 교체할 수 있다. 기존 사건은 사건이 시작될 당시의 세 사람에게 고정되어 중간 교체로 심의 결과가 흔들리지 않게 한다.")
    add_subheading(doc, "학교가 운영 전에 지정해야 할 책임자")
    add_bullets(
        doc,
        [
            "서비스 최종 책임자와 운영 관리자",
            "정식 안건을 배정하고 답변할 교사 또는 부서",
            "비공개 개인 고충·안전 제보의 실제 수신자와 대체 담당자",
            "개인정보·보안 사고 대응 연락망",
            "학생회 임기 변경과 졸업·전학·퇴직 계정 처리 담당자",
        ],
        bullet_num,
    )

    # 4. Service structure
    section_heading(
        doc,
        "04  SERVICE MODEL",
        "서비스는 세 개의 길로 나뉜다",
        "공개 제안, 비공개 고충, 학교 재정·행정 공개는 목적과 열람 범위가 다르므로 한 게시판에 섞지 않는다.",
    )
    add_subheading(doc, "A. 공개 학교 개선 제안")
    add_body(doc, "수업 강화·개편, 기숙사 생활, 시설, 학생자치 제도, 학교 운영방식, 정보 공개 요청처럼 공동체가 함께 논의할 수 있는 주제를 다룬다. 로그인한 학생에게 공개되며 50명 동의를 받으면 정식 안건으로 승격된다.")
    add_subheading(doc, "B. 비공개 개인 고충·안전 제보")
    add_body(doc, "학교폭력, 자해 위험, 개인정보 노출, 성희롱, 차별, 개인 권리 침해처럼 다수결이나 공개 토론에 맡기면 안 되는 사안을 위한 별도 경로다. 50명 동의 없이 지정 담당자에게 전달하고, 일반 학생과 무관한 교사는 열람할 수 없다.")
    add_subheading(doc, "C. 학교 재정·행정 공개")
    add_body(doc, "학교회계 예산·결산, 발전기금, 학교가 공개하기로 한 행사·사업·강사 초청·계약·구매 집행 내역을 출처와 함께 보여준다. 공식 공개 API, 검증된 내부 연동, 권한 있는 교사의 직접 등록 순서로 데이터를 받는다.")
    add_callout(doc, "작성 화면의 핵심", "첫 화면에서 ‘모두와 논의할 개선 제안’과 ‘담당자에게만 보내는 개인 고충·안전 제보’를 명확히 나눈다. 두 경로를 한 폼에 넣어 학생이 실수로 민감한 내용을 공개하지 않게 한다.", color=RED, fill="FCEEEE")

    # 5. Proposal journey
    section_heading(
        doc,
        "05  CORE JOURNEY",
        "제안이 공식 답변과 실행으로 이어지는 과정",
        "사용자는 지금 무엇이 일어나고 있는지, 다음 단계는 무엇인지, 누가 답해야 하는지를 언제든 확인할 수 있어야 한다.",
    )
    add_numbers(
        doc,
        [
            "학생이 로그인한 뒤 익명 공개 또는 실명 공개를 선택하고 제안을 작성한다.",
            "제안은 학생 피드에 공개되고, 작성자는 자동으로 첫 번째 동의자가 된다.",
            "활성 학생 계정은 같은 제안에 한 번만 동의할 수 있다. 동의자 명단은 공개하지 않는다.",
            "50명 미만에서는 본인 동의를 철회할 수 있다. 제안은 학생에게 보이지만 일반 교사의 업무함에는 들어가지 않는다.",
            "50명에 도달하면 제안은 정확히 한 번 정식 안건으로 승격되고 승격 시각과 동의 수가 기록된다.",
            "교사가 안건을 검토하고 공식 답변, 결정 사유, 후속 계획을 등록한다.",
            "안건은 채택·보류·반려 이후에도 실행 중·완료까지 이력을 남긴다.",
        ],
        decimal_num,
    )
    add_subheading(doc, "상태는 두 축으로 관리한다")
    add_table(
        doc,
        ["구분", "상태 예시", "왜 분리하는가"],
        [
            ["업무 처리", "동의 모집 → 정식 안건 → 검토 중 → 채택/보류/반려 → 실행 중 → 완료", "공개 여부와 무관하게 학교의 처리 과정을 정확히 표현"],
            ["공개 상태", "공개, 공개 제한, 심의에 의한 숨김", "문제 글을 숨겨도 업무 기록과 심의 증거는 보존"],
        ],
        [1.05, 3.0, 2.45],
    )
    add_subheading(doc, "교사의 공식 답변에 필요한 정보")
    add_bullets(
        doc,
        [
            "결정 상태와 공식 답변 내용",
            "채택·보류·반려 사유",
            "답변한 교사 또는 담당 부서",
            "후속 계획과 예상 일정(해당하는 경우)",
            "답변 시각과 이후 상태 변경 이력",
        ],
        bullet_num,
    )
    add_callout(doc, "운영정책 필요", "공식 답변 기한과 동의 모집 만료 기간은 아직 정하지 않았다. 학교가 감당할 수 있는 업무량과 학생의 기대를 함께 고려해 정한 뒤 운영 설정으로 관리한다.", color=AMBER, fill="FFF7E6")

    # 6. 50 supporters
    section_heading(
        doc,
        "06  SUPPORT RULE",
        "50명 동의제는 어떻게 공정성을 지키는가",
        "50명은 단순한 숫자가 아니라, 정식 안건의 대표성과 교사의 업무 부담을 함께 조절하는 운영 장치다.",
    )
    add_subheading(doc, "동의의 기본 규칙")
    add_bullets(
        doc,
        [
            "현재 활성화된 학생 계정만 동의할 수 있다.",
            "학생회장과 부회장도 일반 학생과 동일한 한 표를 가진다.",
            "교사 또는 슈퍼 어드민 역할만 가진 계정은 학생 동의를 행사할 수 없다.",
            "작성자의 자동 동의 1표를 포함한다.",
            "50명 도달 이후 동의를 철회해도 다시 일반 제안으로 내려가지 않는다.",
            "화면에는 총 동의 수만 표시하고 동의자의 이름과 내부 식별자를 공개하지 않는다.",
        ],
        bullet_num,
    )
    add_subheading(doc, "중복 동의와 계정 탈취를 구분한다")
    add_body(doc, "같은 계정이 빠르게 여러 요청을 보내는 중복 동의는 데이터베이스의 고유 제약으로 막는다. 시스템은 학번이나 로그인 ID 대신 공개되지 않는 내부 무작위 식별자를 사용한다. 단순 암호화나 Base64 변환은 중복 방지 장치가 아니며, 최종 기준은 ‘제안 1건과 학생 계정 1개당 동의 기록 1개’라는 데이터베이스 규칙이다.")
    add_body(doc, "한 사람이 탈취한 여러 학생 계정으로 동의하는 문제는 다른 종류의 위협이다. 서로 다른 가입 코드, 안전한 비밀번호 저장, 로그인 실패 제한, 이상 로그인 기록, 계정 잠금, 세션 무효화로 계정 탈취 자체를 어렵게 해야 한다.")
    add_callout(doc, "개인정보 기준", "동의자 목록은 제안 작성자, 일반 학생, 일반 교사에게 공개하지 않는다. 운영상 꼭 필요한 경우에도 최소 권한과 감사 기록을 적용한다.")

    # 7. Anonymity and moderation
    section_heading(
        doc,
        "07  ANONYMITY & DUE PROCESS",
        "익명성을 보호하면서도 남용에 대응하는 방법",
        "익명은 ‘아무도 책임지지 않는다’는 뜻이 아니라, 정당한 절차 없이는 권한자도 신원을 볼 수 없다는 약속이다.",
    )
    add_subheading(doc, "익명 제안의 공개 범위")
    add_body(doc, "익명 공개를 선택하면 다른 학생과 일반 교사 화면에 작성자의 이름이 표시되지 않는다. 시스템은 작성자가 본인 글을 수정·확인하고, 심각한 남용에 법과 학교 규정에 따라 대응할 수 있도록 작성자 연결 정보를 보호된 영역에 보관한다. 이 연결 정보는 일반 검색, 로그, 오류 메시지, Swagger 예시에 포함하지 않는다.")
    add_subheading(doc, "사전 검열은 하지 않는다")
    add_body(doc, "작성 단계에서 욕설·비판·부정적 표현을 AI 또는 금칙어 목록으로 자동 차단하지 않는다. 표현의 맥락을 기계가 임의로 판단해 학생의 의견을 막지 않기 위해서다. 다만 스크립트 실행이나 악성 HTML 같은 보안 공격은 입력과 출력 처리로 차단하며, 공개된 글에 대해서는 신고 절차를 제공한다.")
    add_subheading(doc, "공개 제한과 신원 확인은 별개다")
    add_numbers(
        doc,
        [
            "신고가 접수된다. 신고만으로 글을 자동 삭제하거나 신원을 공개하지 않는다.",
            "공개 제한이 필요한지 별도 사건을 만든다. 학생부장·학생회장·부회장이 각각 사유와 함께 판단한다.",
            "세 사람 모두 승인하면 글은 일반 사용자에게 숨겨진다. 원문과 결정 기록은 이의 제기와 감사를 위해 보존한다.",
            "신원 확인이 필요하다면 다시 별도의 사건과 투표를 진행한다. 공개 제한 승인만으로 신원을 볼 수 없다.",
            "신원 확인도 세 사람 모두 승인해야 한다. 학생회장과 부회장은 끝까지 이름과 학번을 보지 않는다.",
            "학생부장만 최근 재인증 후 사유를 남기고 실제 신원을 열람한다. 열람 시각과 열람자가 감사 기록에 남는다.",
        ],
        decimal_num,
    )
    add_callout(doc, "권한 집중 방지", "슈퍼 어드민은 심의위원을 임명하고 계정을 관리할 수 있지만, 그 권한만으로 익명 작성자의 신원을 열람할 수 없다.", color=RED, fill="FCEEEE")

    # 8. Finance disclosure
    section_heading(
        doc,
        "08  TRANSPARENCY",
        "학교 재정·행정 정보를 이해 가능한 형태로 공개한다",
        "원본 문서를 단순히 모아 두는 것이 아니라, 학생이 ‘무엇에 왜 얼마가 쓰였는지’를 출처와 함께 이해할 수 있게 한다.",
    )
    add_subheading(doc, "공개 대상의 예")
    add_bullets(
        doc,
        [
            "학교회계 예산과 결산",
            "학교발전기금 등 공식 공시정보",
            "학생에게 공개하기로 한 행사·사업·강사 초청·계약·구매 집행 내역",
            "원본 문서 또는 공식 출처 링크",
        ],
        bullet_num,
    )
    add_subheading(doc, "데이터를 가져오는 우선순위")
    add_numbers(
        doc,
        [
            "공식 공개 API: 제공 범위와 이용 조건이 확인된 항목만 연동한다.",
            "학교의 검증된 내부 데이터: 권한과 공개 가능 범위를 확인한 뒤 연결한다.",
            "권한 있는 교사의 직접 등록: 외부 연동이 없어도 검증된 자료를 공개할 수 있게 한다.",
        ],
        decimal_num,
    )
    add_body(doc, "표시 항목에는 사업명, 목적, 예산과 실제 집행액의 구분, 금액, 기준일, 분류, 공개 가능한 상대방 정보, 설명, 출처 URL, 등록자·검증자, 최종 수정 시각이 포함된다. 예산과 실제 지출을 같은 값처럼 보이게 하지 않으며 개인정보와 비공개 대상은 게시 전 담당자가 검토한다.")
    add_callout(doc, "금지", "API 키나 정확한 제공 항목이 준비되지 않았을 때 예시 데이터를 실제 정보처럼 만들거나, 비공식 페이지를 무단 수집해 운영 데이터로 사용하지 않는다.", color=RED, fill="FCEEEE")

    # 9. UX
    section_heading(
        doc,
        "09  EXPERIENCE",
        "사람이 망설이지 않는 화면",
        "‘AI가 만든 것 같은’ 장식적 대시보드 대신, 정보와 행동의 우선순위가 자연스럽게 읽히는 절제된 경험을 지향한다.",
    )
    add_subheading(doc, "시각 방향")
    add_bullets(
        doc,
        [
            "Dieter Rams의 ‘Less, but better’와 스칸디나비안 디자인의 절제를 따른다.",
            "흰색과 중립색을 중심으로 하고, 상태와 핵심 행동에만 제한된 강조색을 사용한다.",
            "카드, 그림자, 둥근 모서리, 그라데이션, 글로우, 유리 효과를 반복하지 않는다.",
            "타이포그래피, 여백, 정렬, 대비로 위계를 만든다.",
            "상태를 색상만으로 표현하지 않고 텍스트와 아이콘을 함께 사용한다.",
            "모바일과 데스크톱 모두에서 제안 작성·동의·상태 확인이 짧아야 한다.",
            "키보드 탐색, 명확한 포커스, 입력 라벨, 오류 설명을 기본 품질로 본다.",
        ],
        bullet_num,
    )
    add_subheading(doc, "MVP 핵심 화면")
    add_table(
        doc,
        ["화면", "사용자가 얻어야 하는 답", "핵심 행동"],
        [
            ["로그인·활성화", "내 계정이 안전하게 준비됐는가?", "일회용 코드 사용, 비밀번호 설정, 로그인"],
            ["제안 피드", "지금 어떤 문제가 논의되고 있는가?", "검색·정렬, 제안 읽기, 동의"],
            ["작성 경로 선택", "이 내용은 공개 제안인가, 비공개 고충인가?", "두 경로 중 안전한 선택"],
            ["제안 상세", "몇 명이 동의했고 학교는 어디까지 처리했는가?", "동의·철회, 신고, 공식 답변 확인"],
            ["정식 안건", "학교가 답해야 할 안건은 무엇인가?", "교사 검토, 답변, 상태 변경"],
            ["심의", "무엇을 왜 제한하거나 확인하려는가?", "승인·거부와 사유 기록"],
            ["재정·행정 공개", "무엇에 얼마가 쓰였고 출처는 무엇인가?", "내역 탐색, 원문 확인"],
        ],
        [1.25, 3.15, 2.10],
    )
    add_callout(doc, "첫 화면 우선순위", "브랜드 문구보다 ‘제안 보기’, ‘제안 작성’, ‘비공개 도움 요청’, ‘정식 안건 진행 상황’을 먼저 보여준다.")

    # 10. Security
    section_heading(
        doc,
        "10  SECURITY & PRIVACY",
        "숙련된 학생도 사용하는 환경을 전제로 한다",
        "이 서비스는 인증된 계정이 공격에 사용될 수 있다고 가정한다. 화면을 숨기는 수준이 아니라 서버와 데이터베이스가 권한과 불변조건을 강제해야 한다.",
    )
    add_subheading(doc, "가장 중요한 보호 자산")
    add_bullets(
        doc,
        [
            "익명 작성자와 실제 계정의 연결 정보",
            "비공개 개인 고충·안전 제보 내용",
            "학생·교사 계정과 세션",
            "동의 기록과 50명 승격 결과",
            "심의 투표와 신원 열람 기록",
            "학교 재정·행정 자료의 출처와 변경 이력",
        ],
        bullet_num,
    )
    add_subheading(doc, "필수 방어선")
    add_table(
        doc,
        ["위협", "설계 대응", "확인 방법"],
        [
            ["공용 초기 비밀번호 악용", "사용자별 10~12자 이상 일회용 무작위 가입 코드, 원문 미저장, 사용·만료 후 폐기", "활성화·재사용·만료 통합 테스트"],
            ["비밀번호 탈취", "Argon2id 해시, 로그인 실패 제한, 재설정 코드, 중요 변경 후 세션 폐기", "인증 보안 테스트와 로그 점검"],
            ["권한 상승·IDOR", "모든 객체 권한을 서버에서 재검증, 내부 Entity를 직접 응답하지 않음", "역할별 음성 테스트"],
            ["중복 동의·동시성", "DB 고유 제약과 트랜잭션으로 1인 1표 및 정확히 한 번 승격", "병렬 요청 통합 테스트"],
            ["XSS·CSRF", "안전한 출력, 원문 HTML 금지, CSRF 보호, 보안 쿠키·헤더", "자동 검사와 수동 브라우저 검증"],
            ["익명 신원 노출", "데이터 분리, 일반 API·로그·문서에서 제외, 만장일치와 재인증", "응답 스키마·로그·권한 테스트"],
            ["관리자 오남용", "최소 권한, 역할·심의·열람 감사, 단독 신원 열람 차단", "감사 로그와 운영 점검"],
        ],
        [1.25, 3.45, 1.80],
    )
    add_subheading(doc, "출시 전 검증")
    add_body(doc, "실제 학생 개인정보가 없는 별도 시험 환경을 만들고, 학교의 보안 역량이 높은 학생들이 운영 서버가 아닌 그 환경에서 권한 상승, 신원 노출, 중복 투표, 심의 우회, 세션 공격을 검증한다. 발견된 취약점은 수정 후 같은 시나리오로 재검증하며, 공개 저장소에는 세부 취약점을 수정 전에 노출하지 않는다.")
    add_callout(doc, "출시 기준", "HTTPS, 백업·복구, 비밀정보 관리, 감사 로그, 개인정보 보존·파기 기준, 보안 연락망이 준비되지 않으면 실제 학생 데이터로 운영하지 않는다.", color=RED, fill="FCEEEE")

    # 11. Technical architecture
    section_heading(
        doc,
        "11  DELIVERY ARCHITECTURE",
        "장기 유지보수를 위한 기술 구성",
        "학교에서 가르치는 기술, 몇 년 뒤 다른 학생과 교사가 이어받을 가능성, 보안과 문서화를 함께 고려해 구성한다.",
    )
    add_table(
        doc,
        ["영역", "선택", "이유"],
        [
            ["프런트엔드", "Next.js + TypeScript", "학생·교사 화면을 한 코드베이스에서 관리하고 타입 안전한 API 연동"],
            ["백엔드", "Spring Boot + Java", "명확한 계층과 보안 생태계, 학교 교육과 장기 인수인계 적합성"],
            ["데이터베이스", "MySQL InnoDB", "예상 동시 사용자 규모에 충분하며 트랜잭션·고유 제약 활용"],
            ["세션", "서버 세션을 MySQL에 저장", "MVP에서 Redis 운영 복잡도를 추가하지 않음"],
            ["마이그레이션", "Flyway", "빈 데이터베이스부터 재현 가능한 변경 이력"],
            ["API 문서", "OpenAPI 3 + Swagger UI", "프런트엔드 계약과 신규 개발자 온보딩"],
        ],
        [1.15, 2.2, 3.15],
    )
    add_subheading(doc, "백엔드 구조 원칙")
    add_body(doc, "기능을 Controller·Service·Repository·Entity·DTO로만 한꺼번에 나누는 대신, 인증·사용자·보직·제안·동의·심의·신원·비공개 고충·재정 공개·감사처럼 도메인별로 묶는다. 각 도메인 안에서 Controller, Service, Repository, Entity, Request DTO, Response DTO, Mapper, Exception을 구분한다.")
    add_bullets(
        doc,
        [
            "요청 DTO와 응답 DTO를 같은 클래스로 재사용하지 않는다.",
            "JPA Entity를 API 응답으로 직접 반환하지 않는다.",
            "권한·상태 전이·트랜잭션은 Service가 책임지고, Repository에는 정책 판단을 넣지 않는다.",
            "공통 오류 응답과 전역 예외 처리를 두어 SQL·스택 트레이스·내부 식별자를 외부에 노출하지 않는다.",
            "초기 설계 단계에서 ERD, 권한표, 상태 전이, 위협 모델을 먼저 문서화한다.",
        ],
        bullet_num,
    )
    add_subheading(doc, "Swagger와 README는 완료 조건이다")
    add_body(doc, "Swagger는 모든 주요 흐름을 실제로 시험할 수 있을 정도의 요청·응답 예시, 역할, 오류, 상태 충돌을 포함해야 한다. 생성된 OpenAPI 명세에서 프런트엔드 타입 또는 클라이언트를 만들어 양쪽 계약이 어긋나면 빌드에서 발견되게 한다. README는 저장소를 처음 본 개발자가 환경 준비, 실행, 테스트, 마이그레이션, Swagger 접속, 보안 주의사항을 따라 바로 작업에 들어갈 수 있게 작성한다.")

    # 12. Open source and git
    section_heading(
        doc,
        "12  OPEN SOURCE & GIT",
        "공개할 코드와 지켜야 할 운영 경계를 분리한다",
        "오픈 소스는 코드 공개가 아니라, 다른 사람이 이해하고 기여하고 안전하게 취약점을 알릴 수 있는 운영 체계를 함께 만드는 일이다.",
    )
    add_subheading(doc, "권장 공개 방식")
    add_body(doc, "코드는 Apache License 2.0으로 공개하는 방향을 권장한다. 다만 저작권 주체, 학교명·교표·로고 사용 권한, GitHub 조직 소유자, 장기 유지관리자를 학교와 먼저 확인한다. 애플리케이션 코드와 예제 설정은 공개할 수 있지만 실제 학생·교사 정보, 운영 데이터베이스, 로그, 익명 신원 연결정보, API 키와 서버 비밀값은 저장소와 Git 기록에 절대 포함하지 않는다.")
    add_subheading(doc, "공개 저장소에 필요한 문서")
    add_bullets(
        doc,
        [
            "LICENSE와 정확한 NOTICE",
            "README, CONTRIBUTING, CODE_OF_CONDUCT, GOVERNANCE, SECURITY",
            "Issue·Pull Request 템플릿과 CODEOWNERS",
            "예제 환경설정과 개인정보 없는 데모 데이터",
            "비공개 취약점 신고 경로와 보안 지원 범위",
        ],
        bullet_num,
    )
    add_subheading(doc, "브랜치 전략: 단계에 맞게 바꾼다")
    add_table(
        doc,
        ["단계", "브랜치", "운영 방식"],
        [
            ["MVP", "보호된 main + 짧은 feat/*, fix/*", "작은 PR, 필수 CI, 코드 리뷰, squash merge. develop·release·hotfix는 만들지 않음"],
            ["성숙한 공개 운영", "main + develop + feature/* + release/* + hotfix/*", "안정 릴리스, 외부 기여, 병렬 개발과 긴급 패치가 실제로 필요해질 때 경량 GitFlow로 전환"],
            ["비공개 보안 수정", "GitHub Security Advisory의 private fork", "패치와 공지 시점이 준비되기 전 공개 브랜치에 취약점 세부정보를 올리지 않음"],
        ],
        [1.15, 2.25, 3.10],
    )
    add_callout(doc, "GitFlow 전환 기준", "안정된 운영 릴리스와 외부 기여자가 생기고, 릴리스 준비와 다음 개발을 병렬로 하거나 운영 버전에 긴급 패치가 필요한 시점에 전환한다. 브랜치 수만으로 성숙도를 판단하지 않는다.")

    # 13. MVP and roadmap
    section_heading(
        doc,
        "13  MVP & ROADMAP",
        "첫 버전은 신뢰의 핵심부터 만든다",
        "기능을 많이 넣기보다 로그인, 1인 1표, 익명성, 공식 답변, 심의 절차가 실제로 믿을 수 있게 작동하는 것을 MVP의 성공으로 본다.",
    )
    add_subheading(doc, "MVP에 포함")
    add_bullets(
        doc,
        [
            "자체 계정 생성, 사용자별 일회용 가입 코드, 로그인·로그아웃·비밀번호 재설정",
            "학생·교사·슈퍼 어드민 역할과 임기 기반 3개 보직",
            "익명·실명 공개 제안, 학생 피드, 상세 보기, 본인 작성·동의 내역",
            "1인 1표, 동의 철회, 50명 승격, 정식 안건 목록",
            "교사의 공식 답변과 처리 상태·이력",
            "신고, 공개 제한 심의, 별도 신원 확인 심의, 학생부장 단독 열람",
            "비공개 개인 고충·안전 제보의 분리된 경로와 권한",
            "학교 재정·행정 공개의 수동 등록과 외부 API 어댑터 경계",
            "감사 로그, 속도 제한, 보안 설정, Swagger, README, 핵심 자동 테스트",
        ],
        bullet_num,
    )
    add_subheading(doc, "MVP에서 제외")
    add_bullets(
        doc,
        [
            "첨부파일 업로드",
            "Google/OIDC 로그인",
            "AI 욕설 판별이나 자동 검열",
            "Redis와 불필요한 분산 시스템",
            "여러 학교가 한 서버를 공유하는 멀티테넌시",
            "공식 승인 전 실제 학생 데이터로 외부 공개 운영",
            "확정되지 않은 학교 내부 공문 시스템 연동",
        ],
        bullet_num,
    )
    add_subheading(doc, "단계별 추진")
    add_table(
        doc,
        ["단계", "목표", "주요 산출물", "통과 조건"],
        [
            ["1. 합의", "운영 책임과 정책 확정", "권한표, 상태 전이, 보존정책, 답변 기준", "학교 담당자와 학생 대표의 승인"],
            ["2. 설계", "권한·데이터·위협 경계 고정", "ERD, API 초안, 위협 모델, 화면 흐름", "모호한 항목에 명시적 결정"],
            ["3. MVP", "핵심 흐름 구현", "실행 가능한 앱, Swagger, README, 테스트", "완료 조건과 자동 테스트 통과"],
            ["4. 검증", "보안·사용성 확인", "별도 시험 환경, 취약점 보고, 사용자 테스트", "수정 후 재검증, 치명적 이슈 없음"],
            ["5. 시범 운영", "작은 범위에서 운영 검증", "운영 매뉴얼, 문의·사고 대응, 지표", "책임자 승인과 회고"],
            ["6. 공개", "안전한 오픈 소스 전환", "라이선스·기여·보안 문서, 공개 릴리스", "비밀·개인정보·Git 기록 검사 완료"],
        ],
        [0.75, 1.45, 2.55, 1.75],
    )

    # 14. Measures and risks
    section_heading(
        doc,
        "14  MEASUREMENT & RISK",
        "무엇을 성공으로 볼 것인가",
        "게시글 수가 많다는 사실만으로 성공이라 할 수 없다. 참여, 응답, 실행, 신뢰, 안전을 함께 본다.",
    )
    add_subheading(doc, "측정할 지표")
    add_table(
        doc,
        ["관점", "지표 예시", "읽는 방법"],
        [
            ["참여", "계정 활성화율, 월간 참여 학생, 제안 작성·동의 학생 비율", "소수 사용자만 반복 사용하지 않는지 확인"],
            ["대표성", "50명 도달 비율, 학년별 참여 분포", "문턱이 지나치게 높거나 낮지 않은지 확인"],
            ["응답", "정식 안건 첫 답변 시간, 미답변 안건 수", "학교 업무가 실제로 이어지는지 확인"],
            ["실행", "채택률, 실행 중·완료 전환율, 평균 완료 시간", "답변에서 끝나지 않는지 확인"],
            ["신뢰", "학생·교사 만족도, 상태 설명 이해도, 재이용 의향", "절차가 공정하다고 느끼는지 확인"],
            ["안전", "계정 침해, 권한 우회, 개인정보 노출, 심의 건수", "사고 0을 목표로 하되 작은 징후도 추적"],
        ],
        [1.0, 2.9, 2.6],
    )
    add_body(doc, "구체적인 목표값은 시범 운영 전 기준선을 확인하고 학생·교사와 합의해 정한다. 숫자를 먼저 정해 참여를 강요하거나, 신고 건수가 적다는 이유만으로 안전하다고 판단하지 않는다.")
    add_subheading(doc, "주요 위험과 대응")
    add_table(
        doc,
        ["위험", "가능한 결과", "대응 방향"],
        [
            ["익명성 악용", "비방, 갈등, 피해", "사전 검열 대신 신고·만장일치 심의·기록·교육"],
            ["계정 공유·탈취", "대리 동의, 신원 도용", "개별 가입 코드, 안전한 비밀번호, 로그인 제한, 이상 징후 대응"],
            ["50명 미만 소수자 문제", "중요한 고충이 묻힘", "개인 고충·안전 제보를 별도 경로로 처리"],
            ["교사 업무 증가", "답변 지연, 형식적 처리", "50명 승격, 담당 배정, 답변 기한·상태 기준 합의"],
            ["재정 정보 오해", "예산과 지출 혼동, 불신", "용어 설명, 출처, 기준일, 검증자, 예산·집행 구분"],
            ["관리자 오남용", "신원 노출, 기록 조작", "최소 권한, 심의 분리, 재인증, 변경 불가능한 감사 이력"],
            ["졸업 후 유지 중단", "보안 패치·운영 공백", "학교 소유권·관리자 2인 이상·README·운영 매뉴얼·인수인계"],
            ["공개 저장소 유출", "개인정보·비밀값 노출", "공개 전 전체 Git 기록 검사, secret scanning, 예제 데이터 분리"],
        ],
        [1.35, 2.0, 3.15],
    )

    # 15. Decision register
    section_heading(
        doc,
        "15  DECISION REGISTER",
        "확정된 것과 아직 결정할 것",
        "좋은 기획서는 빈칸을 숨기지 않는다. 다음 항목은 구현 전에 책임 있는 사람이 결정해야 한다.",
    )
    add_subheading(doc, "확정")
    add_bullets(
        doc,
        [
            "모든 사용자는 자체 계정으로 로그인한다.",
            "공통 초기 비밀번호를 사용하지 않고 사용자별 일회용 가입 코드를 발급한다.",
            "학생은 익명 공개와 실명 공개 중 선택한다.",
            "공개 제안은 활성 학생 50명 동의 후 정식 안건이 된다.",
            "동일 계정의 중복 동의는 데이터베이스 고유 제약으로 차단한다.",
            "신원 확인 심의는 학생부장·학생회장·학생부회장 만장일치다.",
            "실제 신원은 학생부장만 재인증 후 열람할 수 있다.",
            "MVP 기술은 Next.js, Spring Boot, MySQL이며 Redis는 사용하지 않는다.",
            "Swagger와 상세 README는 핵심 산출물이다.",
            "MVP는 단순 Feature Branch 방식, 성숙한 공개 운영은 경량 GitFlow를 검토한다.",
        ],
        bullet_num,
    )
    add_subheading(doc, "학교 및 프로젝트 책임자가 결정해야 할 항목")
    add_table(
        doc,
        ["결정 항목", "왜 필요한가", "결정 주체 제안"],
        [
            ["서비스 정식 명칭·로고", "화면·도메인·공개 저장소의 일관성", "학교 담당자 + 학생 대표"],
            ["저작권 주체와 학교 자산 사용", "Apache-2.0 공개와 상표·로고 사용 범위", "학교 관리자 + 프로젝트 책임자"],
            ["공식 답변 기한", "학생 기대와 교사 업무 기준", "학교 운영진 + 교사 대표 + 학생회"],
            ["동의 모집 만료 기간", "오래된 제안 관리와 공정성", "교사 대표 + 학생회"],
            ["정식 안건 배정 방식", "누가 언제 답할지 명확화", "학교 운영진"],
            ["비공개 고충 수신자", "민감한 사안의 안전한 처리", "학생부·상담·학교 관리자"],
            ["이의 제기·재심", "오판 또는 상황 변화 대응", "학교 규정 담당자 + 학생 대표"],
            ["데이터 보존·파기", "개인정보 보호와 감사의 균형", "학교 개인정보 담당자"],
            ["졸업·전학·퇴직 계정", "권한 잔존과 데이터 소유 문제 방지", "학교 운영 관리자"],
            ["호스팅·도메인·백업", "운영 책임과 사고 복구", "학교 IT 담당자 + 개발팀"],
            ["외부 API 제공 항목·키", "실제 연동 범위와 이용 조건", "재정 담당자 + 개발팀"],
            ["Google/OIDC·첨부파일 시점", "추가 보안·운영 비용 판단", "제품 책임자 + 개발팀"],
        ],
        [1.75, 2.8, 1.95],
    )
    add_callout(doc, "다음 회의의 목표", "위 미확정 항목 전체를 한 번에 결정할 필요는 없다. MVP 설계에 직접 영향을 주는 답변 기한, 동의 만료, 담당자, 보존정책, 호스팅 책임을 먼저 확정한다.", color=GREEN, fill="EDF7F2")

    # Appendix
    section_heading(
        doc,
        "APPENDIX",
        "용어와 대표 시나리오",
        "개발자뿐 아니라 학생과 교사가 같은 단어를 같은 의미로 사용하기 위한 참고 자료다.",
    )
    add_subheading(doc, "대표 시나리오")
    add_small_label(doc, "시나리오 A · 기숙사 개선 제안")
    add_body(doc, "학생 A가 ‘기숙사 자습 공간의 운영 시간을 조정해 달라’는 제안을 익명으로 작성한다. 학생 피드에서 49명이 추가로 동의해 총 50명이 되면 정식 안건으로 승격된다. 담당 교사는 현황과 제약을 확인한 뒤 ‘시범 운영 채택’과 일정을 답변한다. 이후 상태는 실행 중, 시범 결과 확인, 완료 순으로 기록된다.")
    add_small_label(doc, "시나리오 B · 비공개 안전 제보")
    add_body(doc, "학생 B가 다른 학생의 개인정보 노출과 관련된 긴급한 내용을 전달하려 한다. 공개 제안이 아니라 비공개 개인 고충·안전 제보를 선택한다. 이 내용은 50명 동의를 받지 않으며 지정 담당자만 확인한다. 일반 학생 피드나 정식 안건 목록에는 나타나지 않는다.")
    add_small_label(doc, "시나리오 C · 익명 신원 확인 요청")
    add_body(doc, "신고된 익명 글에 대해 공개 제한과 신원 확인이 각각 필요한지 검토한다. 학생부장, 학생회장, 부회장은 두 사건을 따로 판단한다. 공개 제한이 만장일치로 승인돼도 신원은 자동 공개되지 않는다. 신원 확인 사건까지 만장일치일 때 학생부장만 재인증 후 확인하며, 학생 임원 화면에는 실제 신원이 표시되지 않는다.")
    add_subheading(doc, "용어")
    add_table(
        doc,
        ["용어", "뜻"],
        [
            ["공개 제안", "로그인한 학생이 함께 읽고 동의할 수 있는 학교 개선 제안"],
            ["정식 안건", "활성 학생 50명의 동의를 받아 교사의 공식 검토 대상으로 승격된 제안"],
            ["비공개 고충", "다수결이나 공개 토론 없이 지정 담당자가 처리해야 하는 개인·안전 사안"],
            ["공개 제한", "DB에서 원문을 지우지 않고 일반 사용자에게 보이지 않게 하는 조치"],
            ["신원 확인", "익명 작성자의 실제 신원을 열람할 수 있는 조건을 판단하는 별도 심의"],
            ["심의위원 스냅샷", "사건 시작 시점의 학생부장·학생회장·부회장을 그 사건 담당자로 고정한 기록"],
            ["감사 로그", "누가 언제 어떤 민감한 행동을 했는지 추적하기 위한 변경·열람 기록"],
            ["MVP", "핵심 가치를 검증할 수 있는 최소 기능 제품. ‘대충 만든 버전’이 아니라 핵심 신뢰가 작동하는 첫 버전"],
        ],
        [1.55, 4.95],
    )
    add_subheading(doc, "연결 문서")
    add_bullets(
        doc,
        [
            "AI 개발 에이전트용 상세 구현 지시서: AI_AGENT_DEVELOPMENT_PROMPT.md",
            "향후 작성할 문서: ERD, 권한 매트릭스, 상태 전이표, 위협 모델, 운영 매뉴얼, 개인정보 처리 기준, 보안 정책",
        ],
        bullet_num,
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    keep_with_next(p)
    r = p.add_run("참고 기준")
    style_run(r, bold=True, size=10, color=BLUE_DARK)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(3)
    style_run(p2.add_run("OWASP Application Security Verification Standard: "), size=8.5, color=MUTED)
    add_hyperlink(p2, "owasp.org/www-project-application-security-verification-standard", "https://owasp.org/www-project-application-security-verification-standard/", color=BLUE)
    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(3)
    style_run(p3.add_run("Apache License 2.0: "), size=8.5, color=MUTED)
    add_hyperlink(p3, "apache.org/licenses/LICENSE-2.0", "https://www.apache.org/licenses/LICENSE-2.0", color=BLUE)

    # Ensure every section shares the same dimensions and header/footer behavior.
    for section in doc.sections:
        section.page_width = PAGE_WIDTH
        section.page_height = PAGE_HEIGHT
        section.top_margin = MARGIN
        section.bottom_margin = MARGIN
        section.left_margin = MARGIN
        section.right_margin = MARGIN
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
